import os
from pathlib import Path
from docx import Document
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph

def replace_text_in_paragraph(paragraph: Paragraph, replacements: dict, found_keys: set):
    """
    Replace the text in the paragraph with the replacements.
    """
    text = "".join(run.text for run in paragraph.runs)
    original_text = text

    for key, val in replacements.items():
        if key in text:
            text = text.replace(key, val)
            found_keys.add(key)

    if text != original_text:
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        # Write the new full text into the first run
        if paragraph.runs:
            paragraph.runs[0].text = text

def replace_placeholders(doc: _Document, replacements: dict):
    """
    Replace the placeholders in the document with the replacements.
    """
    found_keys = set()

    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements, found_keys)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements, found_keys)

    missing = set(replacements.keys()) - found_keys
    if missing:
        raise ValueError(f"Missing placeholders in template: {', '.join(missing)}")

    return doc

def generate_certificates(template_path: str, eligible_attendees: list[dict]):
    """
    Generate the certificates for the eligible attendees.
    The certificates will be saved in the Downloads/certs/docx folder.
    """

    saved_certs = []

    # Set the output directory
    docx_output_dir = str(Path.home() / "Downloads" / "certs" / "docx")
    os.makedirs(docx_output_dir, exist_ok=True)

    # Capitalize the first letter of each word in the first name and uppercase last name
    for attendee in eligible_attendees:
        attendee["名字"] = " ".join(word.capitalize() for word in attendee["名字"].split(" "))
        attendee["姓氏"] = attendee["姓氏"].upper()

    for idx, attendee in enumerate(eligible_attendees):
        try:
            ref_num = str(idx + 1).zfill(2)
            replacements = {
                "<<first_name>>": attendee.get("名字", ""),
                "<<last_name>>": attendee.get("姓氏", ""),
                "<<ref>>": ref_num
            }

            # Load a fresh template
            doc = Document(template_path)

            # Replace placeholders
            replace_placeholders(doc, replacements)

            # Save certificate
            base_name = f"{ref_num}_{attendee.get('名字', 'Unknown')}_{attendee.get('姓氏', 'Unknown')}"
            docx_path = os.path.join(docx_output_dir, f"{base_name}.docx")
            doc.save(docx_path)

            print(f"Saved: {docx_path}")
            saved_certs.append(docx_path)

        except Exception as e:
            print(f"Skipping attendee {attendee.get('電子郵件地址', '')}: {e}")

    return saved_certs



